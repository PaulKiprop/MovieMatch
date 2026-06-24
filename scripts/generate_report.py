"""Generate the MovieMatch project report as a .docx file.

Run:  python scripts/generate_report.py
Output: MovieMatch_Report.docx (in the project root)

The document mirrors REPORT.md and inserts clearly-marked placeholders where
the user should drop in their own screenshots.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "MovieMatch_Report.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x66, 0x66, 0x66)


def add_placeholder(doc, caption):
    """Insert a bordered, shaded box telling the user where to paste an image."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Pt(420)
    # Shade the cell so the placeholder is obvious.
    tcPr = cell._tc.get_or_add_tcPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "EFEFEF")
    tcPr.append(shd)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[  SCREENSHOT PLACEHOLDER  ]\n{caption}\n")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run(f"Figure: {caption}")
    crun.italic = True
    crun.font.size = Pt(9)
    crun.font.color.rgb = GREY
    doc.add_paragraph()


def h(doc, text, level):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(head)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("MovieMatch")
    trun.bold = True
    trun.font.size = Pt(26)
    trun.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run("Intelligent Movie Recommendation System — Project Report")
    srun.font.size = Pt(13)
    srun.font.color.rgb = GREY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mrun = meta.add_run("Module: Computer Programming in Python · GISMA\n"
                        "Application type: Desktop GUI application (Python + Tkinter)")
    mrun.font.size = Pt(10)
    mrun.font.color.rgb = GREY
    doc.add_paragraph()

    # ---- 1. Introduction ----
    h(doc, "1. Introduction", 1)
    para(doc,
         "MovieMatch is a desktop application that helps users discover films "
         "based on their personal tastes and viewing history. Rather than "
         "browsing endless lists, the user builds up a profile simply by rating "
         "and saving the movies they like, and the application responds with "
         "tailored, explainable recommendations.")
    para(doc,
         "The project demonstrates a complete, layered Python application: a live "
         "connection to a real web API, local data persistence, a content-based "
         "recommendation algorithm, data analytics with charts, and a responsive "
         "graphical user interface — all organised into clearly separated modules.")
    para(doc, "The six core capabilities delivered are:", bold=True)
    numbered(doc, [
        "Discover the latest movies automatically on launch (Home page)",
        "Search for movies by title",
        "View movie details (overview, genres, poster, community score)",
        "Save favourites and rate movies (1–5 stars)",
        "Receive recommendations based on taste, with reasons",
        "Analyse watching habits through visual charts",
    ])
    add_placeholder(doc, "Home page showing the latest movies as a poster grid")

    # ---- 2. Implementation Stack ----
    h(doc, "2. Implementation Stack", 1)
    make_table(doc,
               ["Layer", "Technology", "Why it was chosen"],
               [
                   ["Language", "Python 3.14", "Module requirement; rich standard library"],
                   ["GUI", "Tkinter / ttk", "Built into Python — portable, no install needed"],
                   ["Movie data", "TMDB REST API", "Free source of real movie metadata, posters, ratings"],
                   ["HTTP client", "requests", "Clean, reliable API calls"],
                   ["Images", "Pillow (PIL)", "Decodes/resizes TMDB JPEG posters for Tkinter"],
                   ["Charts", "matplotlib", "Embeds analytics charts inside the window"],
                   ["Storage", "SQLite (sqlite3)", "Zero-config local database; standard library"],
               ])
    para(doc,
         "External dependencies are limited to four well-established packages "
         "(requests, Pillow, matplotlib, python-docx), all listed in "
         "requirements.txt.")

    # ---- 3. Architecture ----
    h(doc, "3. Architecture", 1)
    para(doc,
         "The code is organised into a Python package, moviematch, with each "
         "concern isolated in its own sub-module. This separation makes the "
         "system easy to understand, test and extend.")
    code = doc.add_paragraph()
    crun = code.add_run(
        "moviematch/\n"
        "├── config.py                 # API key loading, paths, constants\n"
        "├── models.py                 # the Movie dataclass (shared everywhere)\n"
        "├── api/tmdb_client.py         # TMDB REST wrapper\n"
        "├── data/storage.py            # SQLite persistence layer\n"
        "├── recommender/content_based.py # recommendation algorithm\n"
        "├── analytics/habits.py        # habit aggregation (pure functions)\n"
        "└── gui/                       # Tkinter UI: app + 5 tabs + details window"
    )
    crun.font.name = "Consolas"
    crun.font.size = Pt(9)
    para(doc, "Design principles applied:", bold=True)
    bullets(doc, [
        "Single source of truth — one Movie dataclass is used by every layer, "
        "so all parts of the system speak the same language.",
        "Separation of concerns — networking, storage, logic and presentation "
        "never mix; the analytics module is pure functions, making it directly testable.",
        "Responsiveness — all network calls run on background threads, so the "
        "interface never freezes; results return to the main thread via Tkinter's after().",
        "Graceful degradation — without an API key the app still opens; favourites "
        "and analytics work offline and online features show a clear message.",
    ])

    # ---- 4. Key Features ----
    h(doc, "4. Key Features in Detail", 1)
    h(doc, "4.1 Home (discovery) page", 2)
    para(doc,
         "On startup the app automatically loads the latest cinema releases into a "
         "responsive poster grid. A dropdown switches between Now Playing, Popular, "
         "Top Rated and Upcoming. Each film is a card showing its poster, a gold "
         "rating badge and its title; posters load lazily so the page appears instantly.")
    h(doc, "4.2 Search and details", 2)
    para(doc,
         "The user searches by title; results appear in the same poster grid. "
         "Clicking a poster opens a details window with the full overview, genres, "
         "poster and score, plus controls to rate and favourite the movie.")
    add_placeholder(doc, "Movie details window with poster, rating stars and favourite button")
    h(doc, "4.3 Favourites and ratings", 2)
    para(doc,
         "Favourites and 1–5 star ratings are stored in the local SQLite database "
         "and persist between sessions. Viewing a movie's details is also logged, "
         "building the history used by analytics and recommendations.")
    h(doc, "4.4 Recommendations (content-based)", 2)
    para(doc, "The recommendation engine works in four explainable steps:")
    numbered(doc, [
        "Build a taste profile — each favourited/highly-rated movie contributes a "
        "weighted bag of genres and keywords; 5-star films count more than 3-star, "
        "genres more than keywords, and 1–2 star films are down-weighted.",
        "Gather candidates — fresh movies are pulled from TMDB by the user's top "
        "genres, plus popular titles to broaden the pool.",
        "Score each candidate against the profile using cosine similarity, excluding "
        "anything already seen.",
        "Explain — every recommendation shows a match % and the shared genres/keywords "
        "behind it, so suggestions are transparent rather than a black box.",
    ])
    add_placeholder(doc, "Recommendations tab: poster grid with match % and reasons")
    h(doc, "4.5 Analytics", 2)
    para(doc,
         "The Analytics tab summarises viewing habits and renders three matplotlib "
         "charts embedded in the window: top genres, ratings distribution and movies "
         "by decade, alongside totals such as average rating and most-watched genre.")
    add_placeholder(doc, "Analytics tab with top-genres, ratings and decade charts")

    # ---- 5. Data Persistence ----
    h(doc, "5. Data Persistence", 1)
    para(doc, "A single SQLite database (data_store/moviematch.db) holds four tables:")
    bullets(doc, [
        "movies — a cache of movie metadata, so favourites/analytics work offline",
        "favorites — saved movies",
        "ratings — the user's 1–5 star ratings (constrained at database level)",
        "history — a log of viewed movies",
    ])
    para(doc,
         "Movie-specific tables reference movies(movie_id) as a foreign key, avoiding "
         "duplicated data. All writes use parameterised SQL statements, protecting "
         "against SQL injection and malformed input.")

    # ---- 6. Testing ----
    h(doc, "6. Testing and Quality", 1)
    para(doc,
         "A suite of offline unit tests (tests/test_core.py, run with "
         "python -m unittest) covers the model, storage, recommender and analytics "
         "layers without needing the network or a display. Tests verify that ratings "
         "are validated (1–5 only), that recommendations prefer genre overlap and "
         "exclude already-seen films, and that analytics correctly aggregate genres, "
         "decades and average ratings. Keeping logic free of GUI code is what makes "
         "this isolated testing possible — a direct benefit of the layered architecture.")

    # ---- 7. How to Run ----
    h(doc, "7. How to Run", 1)
    run_code = doc.add_paragraph()
    rrun = run_code.add_run(
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate            # Windows\n"
        "pip install -r requirements.txt\n\n"
        "# Add a free TMDB API key:\n"
        "#   copy .env.example to .env and paste your key\n\n"
        "python main.py"
    )
    rrun.font.name = "Consolas"
    rrun.font.size = Pt(9)
    para(doc,
         "A free API key is obtained from the TMDB website (Settings → API). The "
         "application launches even without one, with online features disabled.")

    # ---- 8. Challenges ----
    h(doc, "8. Challenges and Future Work", 1)
    para(doc, "Challenges addressed:", bold=True)
    bullets(doc, [
        "Keeping the UI responsive during network calls — solved with a background-"
        "thread helper that marshals results back to the main thread.",
        "Displaying JPEG posters in Tkinter (which supports few formats natively) — "
        "solved with Pillow, creating Tk image objects on the main thread only.",
        "Making recommendations explainable — solved by reporting the shared features "
        "behind every suggestion.",
    ])
    para(doc, "Possible extensions:", bold=True)
    bullets(doc, [
        "Add a lightweight collaborative-filtering layer ('users who liked X…')",
        "Support multiple user profiles",
        "Export the analytics charts to an image or PDF report",
        "Cache search results to reduce repeated API calls",
    ])

    # ---- 9. Conclusion ----
    h(doc, "9. Conclusion", 1)
    para(doc,
         "MovieMatch fulfils all six required capabilities in a clean, modular and "
         "well-tested application. It combines a real-world web API, local "
         "persistence, an explainable recommendation algorithm and embedded data "
         "visualisation behind a polished, responsive Tkinter interface — "
         "demonstrating practical application of core Python concepts: object-"
         "oriented design, modular architecture, data handling, concurrency and GUI "
         "development.")

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
